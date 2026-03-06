from __future__ import annotations

import json
from pathlib import Path
from uuid import uuid4

from docx import Document
from openpyxl import Workbook

from app.core.config import get_settings


def build_preview_markdown(
    *,
    project_name: str,
    draft_content: str,
    questions: list[dict],
    warnings: list[dict],
    extraction_summary: dict | None,
) -> str:
    lines = [f"# {project_name} Export Preview", ""]

    if extraction_summary:
        lines.extend(
            [
                "## RFP 요약",
                f"- 사업 개요: {extraction_summary.get('project_summary_text', '').strip() or 'TBD'}",
                f"- 요구사항 수: {extraction_summary.get('requirements_count', 0)}",
                f"- 평가항목 수: {extraction_summary.get('evaluation_count', 0)}",
                "",
            ]
        )

    lines.extend(["## Draft", draft_content.strip(), ""])

    lines.append("## Open Questions")
    if questions:
        lines.extend([f"- [{question['status']}] {question['question_text']}" for question in questions])
    else:
        lines.append("- 없음")
    lines.append("")

    lines.append("## Mapping Warnings")
    if warnings:
        lines.extend([f"- ({warning['type']}) {warning['message']}" for warning in warnings])
    else:
        lines.append("- 없음")
    lines.append("")
    return "\n".join(lines).strip() + "\n"


def _write_docx(path: Path, preview_md: str) -> None:
    document = Document()
    for line in preview_md.splitlines():
        if line.startswith("# "):
            document.add_heading(line.removeprefix("# ").strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line.removeprefix("## ").strip(), level=2)
        else:
            document.add_paragraph(line)
    document.save(path)


def _write_xlsx(path: Path, questions: list[dict], warnings: list[dict]) -> None:
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Summary"
    summary.append(["Metric", "Count"])
    summary.append(["Open Questions", len(questions)])
    summary.append(["Mapping Warnings", len(warnings)])

    question_sheet = workbook.create_sheet("Questions")
    question_sheet.append(["ID", "Status", "Question"])
    for question in questions:
        question_sheet.append([question["id"], question["status"], question["question_text"]])

    warning_sheet = workbook.create_sheet("Warnings")
    warning_sheet.append(["Type", "Message"])
    for warning in warnings:
        warning_sheet.append([warning["type"], warning["message"]])

    workbook.save(path)


def create_export_artifacts(
    *,
    project_id: int,
    project_name: str,
    formats: list[str],
    draft_content: str,
    questions: list[dict],
    warnings: list[dict],
    extraction_summary: dict | None,
) -> tuple[str, str, str]:
    settings = get_settings()
    session_id = uuid4().hex[:12]
    session_dir = settings.export_dir / f"project_{project_id}" / session_id
    session_dir.mkdir(parents=True, exist_ok=True)

    preview_md = build_preview_markdown(
        project_name=project_name,
        draft_content=draft_content,
        questions=questions,
        warnings=warnings,
        extraction_summary=extraction_summary,
    )

    preview_path = session_dir / "preview.md"
    preview_path.write_text(preview_md, encoding="utf-8")

    files: dict[str, str] = {"md": str(preview_path.relative_to(settings.app_data_dir))}

    if "txt" in formats:
        txt_path = session_dir / "preview.txt"
        txt_path.write_text(preview_md, encoding="utf-8")
        files["txt"] = str(txt_path.relative_to(settings.app_data_dir))
    if "docx" in formats:
        docx_path = session_dir / "preview.docx"
        _write_docx(docx_path, preview_md)
        files["docx"] = str(docx_path.relative_to(settings.app_data_dir))
    if "xlsx" in formats:
        xlsx_path = session_dir / "preview.xlsx"
        _write_xlsx(xlsx_path, questions, warnings)
        files["xlsx"] = str(xlsx_path.relative_to(settings.app_data_dir))

    return session_id, str(preview_path.relative_to(settings.app_data_dir)), json.dumps(files)
